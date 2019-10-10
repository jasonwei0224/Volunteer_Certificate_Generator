import docx
from docx.shared import Pt
from typing import List
class WordFile:
    """
    WordFile Class
    """

    def __init__(self, filename: str) -> str:
        """
        Constructor
        Takes a filename

        filename: the filename or path of the file
        content: list of the content
        word_doc: open the document
        """

        self.__filename = filename
        self.__content = []
        self.__word_doc = docx.Document(filename)

    def get_filename(self) -> str:
        """
        Return the file name
        """
        return self.__filename

    def get_file_content(self) -> List[str]:
        """
        Return the file content in a list
        """
        for paragraph in self.__word_doc.paragraphs:
            self.__content.append('' + paragraph.text)
        return self.__content

    def modify_content(self, placeholder: str, replacement: str) -> None:
        """
        Modify the placeholder with the replacement string
        """
        for paragraph in self.__word_doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)

    def replace_string(self, place_holder: str, replacement: str) -> None:
        """
        Replace the string in the content lst
        """
        for index in range(0, len(self.__content)):
            print(self.__content[index])
            self.__content[index] = self.__content[index].replace(place_holder, replacement)

    def save(self, filepath: str) -> None:
        """
        Save the file
        """
        self.__word_doc.save(filepath)


    def set_font(self, name: str, size: int) -> None:
        """
        Set the font to the given name and size
        """
        style = self.__word_doc.styles['Normal']
        font = style.font
        font.name = name
        font.size = Pt(size)
